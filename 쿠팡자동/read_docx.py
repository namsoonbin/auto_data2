import sys
from docx import Document

# UTF-8 출력 설정
sys.stdout.reconfigure(encoding='utf-8')

# 첫 번째 문서 읽기
print("=" * 80)
print("📄 쿠팡 스토어 매출_광고 실적 자동화 웹앱 설계 및 구현.docx")
print("=" * 80)
doc1 = Document('쿠팡 스토어 매출_광고 실적 자동화 웹앱 설계 및 구현.docx')
for para in doc1.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n\n")
print("=" * 80)
print("📄 쿠팡 스토어별 매출_마진 관리 현황과 자동화 구현 제안.docx")
print("=" * 80)
doc2 = Document('쿠팡 스토어별 매출_마진 관리 현황과 자동화 구현 제안.docx')
for para in doc2.paragraphs:
    if para.text.strip():
        print(para.text)
